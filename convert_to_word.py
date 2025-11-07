#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create a new Document
doc = Document()

# Set document title
title = doc.add_heading('Business Requirements Document (BRD)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('KOSH 2.0 - Warehouse Inventory Management Application', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Document info
doc.add_paragraph('Document Version: 1.0')
doc.add_paragraph('Date: October 31, 2025')
doc.add_paragraph('Status: Draft')
doc.add_paragraph()

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'KOSH 2.0 is a warehouse inventory management application designed to replace the legacy Microsoft Access system. '
    'The application manages three distinct inventory types (PCB, Warehouse Goods, and Finished Goods) and streamlines '
    'warehouse operations including receiving, put away, picking, counting, and shortage reporting.'
)

doc.add_heading('Business Objectives', level=2)
objectives = [
    'Replace legacy Access-based inventory system with modern web application',
    'Provide real-time inventory visibility across multiple inventory types',
    'Automate data entry and label generation via API integration',
    'Improve accuracy in warehouse operations',
    'Enable efficient shortage reporting for purchasing decisions'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_page_break()

# 2. Inventory Types
doc.add_heading('2. Inventory Types', level=1)
doc.add_paragraph('The system supports three independent inventory categories:')
inventory_types = [
    'PCB Inventory - Printed Circuit Board management',
    'Warehouse Goods Inventory - General warehouse materials and components',
    'Finished Goods Inventory - Completed products ready for shipment'
]
for inv_type in inventory_types:
    doc.add_paragraph(inv_type, style='List Number')
doc.add_paragraph('Each inventory type maintains independent tracking and menu options.')

# 3. Core Functional Modules
doc.add_heading('3. Core Functional Modules', level=1)

# 3.1 Receipt Entry
doc.add_heading('3.1 Receipt Entry', level=2)
doc.add_paragraph('Data Entry:', style='Heading 3')
receipt_data = [
    'Part Number (manual)',
    'PO Number (manual)',
    'Quantity (manual)',
    'Manufacturer Part Number (MFN) - Auto-generated via API',
    'Date Code (manual)',
    'MSD designation (manual)'
]
for item in receipt_data:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Requirements:', style='Heading 3')
receipt_reqs = [
    'Validate part numbers against master data',
    'Auto-generate MFN by calling external API',
    'Support barcode scanning',
    'Create unique receipt records with timestamps'
]
for req in receipt_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 3.2 Put Away
doc.add_heading('3.2 Put Away', level=2)
putaway_reqs = [
    'Display pending put-away items from receipts',
    'Assign storage locations to received items',
    'Update inventory locations upon completion',
    'Support location barcode scanning'
]
for req in putaway_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 3.3 Pick Operations
doc.add_heading('3.3 Pick Operations', level=2)
doc.add_paragraph('Data Entry:', style='Heading 3')
pick_data = [
    'Work Order number',
    'PCN (scan or manual)',
    'PCN Quantity',
    'Job numbers (e.g., 23806-3, Quarantine, Purge)'
]
for item in pick_data:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Requirements:', style='Heading 3')
pick_reqs = [
    'Validate work orders before allowing picks',
    'Scan PCN barcodes for confirmation',
    'Support partial picks',
    'Update inventory in real-time',
    'Track pick history by user, date, and work order'
]
for req in pick_reqs:
    doc.add_paragraph(req, style='List Bullet')
p = doc.add_paragraph('CRITICAL: Pick data entry area must NOT auto-generate data')
p.runs[0].bold = True

# 3.4 Count Back
doc.add_heading('3.4 Count Back', level=2)
doc.add_paragraph('Data Entry:', style='Heading 3')
countback_data = [
    'PCN (scan or manual)',
    'PCN Quantity',
    'Part Number',
    'PO Number'
]
for item in countback_data:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Requirements:', style='Heading 3')
countback_reqs = [
    'Allow counting back unused materials to inventory',
    'Validate PCN before accepting',
    'Update inventory quantities upon completion',
    'Record transactions with user and timestamp'
]
for req in countback_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 3.5 Cycle Count
doc.add_heading('3.5 Cycle Count', level=2)
doc.add_paragraph('Data Entry:', style='Heading 3')
cycle_data = [
    'PCN (scan or manual)',
    'Part Number',
    'Quantity',
    'PO Number'
]
for item in cycle_data:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Requirements:', style='Heading 3')
cycle_reqs = [
    'Support scheduled and ad-hoc cycle counts',
    'Compare counted vs. system quantities',
    'Flag discrepancies for review',
    'Require supervisor approval for variance adjustments',
    'Maintain complete audit trail'
]
for req in cycle_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_page_break()

# 3.6 PCN Management
doc.add_heading('3.6 PCN (Part Change Notice) Management', level=2)

doc.add_paragraph('Generate PCN:', style='Heading 3')
pcn_gen = [
    'Generate unique PCN identifiers with scannable barcodes',
    'Support creating PCNs for same part when jobs are created consecutively'
]
for item in pcn_gen:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph('CRITICAL: All information must be entered each time for same part')
p.runs[0].bold = True

doc.add_paragraph('PCN History:', style='Heading 3')
pcn_history = [
    'Maintain complete history of each PCN',
    'Track locations where PCN has been used',
    'Track part number changes',
    'Track if PCN has been picked for any jobs',
    'Refresh capability after every entry'
]
for item in pcn_history:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('PCN Status Categories:', style='Heading 3')
pcn_status = ['PCB', 'Partial', 'Completed', 'Shipped', '(Additional categories as needed)']
for item in pcn_status:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Delete Function:', style='Heading 3')
pcn_delete = [
    'Users can delete PCNs they created',
    'Display error message when deletion is not allowed'
]
for item in pcn_delete:
    doc.add_paragraph(item, style='List Bullet')
p = doc.add_paragraph('CRITICAL: System must prevent deletion of next.cfg file')
p.runs[0].bold = True

# 3.7 Stock Management
doc.add_heading('3.7 Stock Management', level=2)
p = doc.add_paragraph('CRITICAL: Must NOT auto-generate data when PCN is scanned or typed')
p.runs[0].bold = True
stock_reqs = [
    'Provide option to print auto-generated label after data inquiry',
    'With Access permission, allow printing of new labels'
]
for req in stock_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 3.8 BOM Loader
doc.add_heading('3.8 BOM Loader', level=2)
bom_reqs = [
    'View current BOM revision',
    'Download current revision',
    'Load ACI-created parts',
    'Validate BOM data before loading'
]
for req in bom_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 3.9 Shortage Reports
doc.add_heading('3.9 Shortage Reports', level=2)
doc.add_paragraph('Report Types:', style='Heading 3')
report_types = [
    'SMT (Surface Mount Technology) Inventory Shortage Report',
    'PTH (Plated Through Hole) Inventory Shortage Report'
]
for item in report_types:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Requirements:', style='Heading 3')
shortage_reqs = [
    'Create shortage reports for each job',
    'Inform Purchasing what to buy and quantity needed',
    'Generate separate reports for SMT and PTH warehouse locations',
    'Support on-demand report creation',
    'Support printing'
]
for req in shortage_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 3.10 Part Number Change
doc.add_heading('3.10 Part Number Change Functionality', level=2)
part_change_reqs = [
    'Support changing part numbers with full audit trail',
    'Maintain linkage between old and new part numbers',
    'Update all related inventory records',
    'Require authorization for changes'
]
for req in part_change_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_page_break()

# 4. User Roles
doc.add_heading('4. User Roles and Permissions', level=1)
roles = [
    ('Warehouse Operator', 'Receive, put away, pick, count back, cycle count, view inventory'),
    ('Warehouse Supervisor', 'All Operator permissions + approve variances, generate shortage reports, access PCN history, print labels'),
    ('Inventory Manager', 'All Supervisor permissions + manage part changes, delete PCNs (restricted), load BOMs, configure system'),
    ('Administrator', 'Full system access, user management, system configuration, audit logs')
]

for role, permissions in roles:
    p = doc.add_paragraph()
    p.add_run(role + ': ').bold = True
    p.add_run(permissions)

# 5. Key Reports
doc.add_heading('5. Key Reports', level=1)
reports = [
    'Inventory Status Report - Current quantities by part, location, and inventory type',
    'Transaction History Report - All inventory movements by date range',
    'Shortage Report (SMT) - Parts needed for SMT operations',
    'Shortage Report (PTH) - Parts needed for PTH operations',
    'PCN History Report - Complete lifecycle of part change notices',
    'Cycle Count Variance Report - Discrepancies during cycle counts',
    'Pick History Report - All picks by work order, date, user'
]
for report in reports:
    doc.add_paragraph(report, style='List Number')

doc.add_paragraph('Report Features:', style='Heading 3')
report_features = [
    'All reports printable and exportable (PDF, Excel)',
    'Support filtering and sorting',
    'Include timestamp and user identification'
]
for feature in report_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# 6. Non-Functional Requirements
doc.add_heading('6. Non-Functional Requirements', level=1)

doc.add_paragraph('Performance:', style='Heading 3')
perf_reqs = [
    'Barcode scanning response time < 1 second',
    'Reports generate within 5 seconds',
    'Support minimum 20 concurrent users',
    '99.5% uptime during business hours'
]
for req in perf_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph('Security:', style='Heading 3')
sec_reqs = [
    'Role-based access control',
    'All transactions logged with user identification',
    'Prevent unauthorized data deletion',
    'Encrypt sensitive data'
]
for req in sec_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph('Data Management:', style='Heading 3')
data_reqs = [
    'Complete audit trail of all transactions',
    'Automated daily backups',
    'Retain transaction history for 7 years',
    'Prevent deletion of critical configuration files'
]
for req in data_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph('Integration:', style='Heading 3')
int_reqs = [
    'Integrate with existing API for MFN auto-generation',
    'Support integration with ERP/Purchasing systems',
    'Export data in standard formats (CSV, Excel, PDF)'
]
for req in int_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 7. Migration Requirements
doc.add_heading('7. Migration Requirements', level=1)
migration_reqs = [
    'Migrate all active inventory data from Access database',
    'Migrate historical transaction data (minimum 2 years)',
    'Migrate complete PCN history',
    'Validate master data (parts, locations, BOMs) during migration',
    'Support parallel operation period with legacy system',
    'Establish data reconciliation process'
]
for req in migration_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_page_break()

# 8. Critical Issues Addressed
doc.add_heading('8. Critical Issues Addressed from Legacy System', level=1)
doc.add_paragraph('Based on testing of previous PCB system, KOSH 2.0 addresses:')
issues = [
    'Bar code length and scannability issues',
    'Inability to use previous items in dropdown menus',
    'Need to refresh after every entry',
    'Inventory dropdown limitations when PCNs created for multiple companies',
    'Delete function errors with next.cfg file',
    'Pick data entry unwanted auto-generation',
    'Stock feature unwanted auto-generation',
    'Lack of separation between inventory types'
]
for issue in issues:
    doc.add_paragraph('✓ ' + issue, style='List Bullet')

# 9. Success Criteria
doc.add_heading('9. Success Criteria', level=1)

doc.add_paragraph('Functional Success:', style='Heading 3')
func_success = [
    'All core modules operational (Receipt, Put Away, Pick, Count, Cycle Count)',
    'All three inventory types independently managed',
    'API integration for MFN generation working',
    'Accurate shortage reports for SMT and PTH'
]
for item in func_success:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('User Acceptance:', style='Heading 3')
user_success = [
    'Warehouse staff trained and proficient',
    'Operations faster than legacy system',
    'Error rates reduced by 50%'
]
for item in user_success:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Technical Success:', style='Heading 3')
tech_success = [
    'All performance requirements met',
    'Zero data loss during migration',
    'All security requirements implemented',
    'Backup and recovery procedures tested'
]
for item in tech_success:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 10. Glossary
doc.add_heading('10. Glossary', level=1)
glossary = [
    ('PCN', 'Part Change Notice: Unique identifier for tracking parts'),
    ('MFN', 'Manufacturer Part Number'),
    ('SMT', 'Surface Mount Technology'),
    ('PTH', 'Plated Through Hole'),
    ('BOM', 'Bill of Materials'),
    ('MSD', 'Moisture Sensitive Device'),
    ('DC', 'Date Code'),
    ('ACI', 'Company-created parts designation')
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term + ': ').bold = True
    p.add_run(definition)

# 11. Approval Table
doc.add_heading('11. Approval Signatures', level=1)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Role'
header_cells[1].text = 'Name'
header_cells[2].text = 'Signature'
header_cells[3].text = 'Date'

roles_for_approval = ['Business Owner', 'Project Sponsor', 'IT Manager', 'Warehouse Manager']
for idx, role in enumerate(roles_for_approval, 1):
    table.rows[idx].cells[0].text = role

# Save the document
doc.save('/home/tony/ACI Invertory/KOSH_2.0_BRD.docx')
print("Word document created successfully: KOSH_2.0_BRD.docx")
